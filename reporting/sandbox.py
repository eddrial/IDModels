'''
Created on 15 Jul 2021

@author: oqb
'''

from docx import Document
import argparse
from apple2p5 import model2 as id
from idanalysis import analysis_functions as af
from idcomponents import parameters
import numpy as np
from time import sleep as sleep
from wradia import wrad_obj as wrd
import radia as rd
import vtk


class Sandbox(object):
    '''
    classdocs
    '''


    def __init__(self, myID, mysolution):
        '''
        Constructor
        '''
        
        self.rep = Document()
        
        
        self.myID = myID
        self.mysolution = mysolution
        
        self.fignum = 1
        self.tabnum = 1
        self.appnum = 1
        
        
        
        
    def publish_report(self,fname, 
                       all = False, 
                       Title = False, 
                       Contents = False, 
                       Introduction = False,
                       Key_Parameters = False,
                       Field_Results = False):
        
        if all == True:
            Title = True
            Contents = True
            Introduction = True
            Key_Parameters = True
            Field_Results = True
        
        if Title == True:
            self.include_title()
            
        if Contents == True:
            self.include_contents()
        
        if Introduction == True:
            self.include_introduction()
        
        if Key_Parameters == True:
            pass
            #self.include_key_parameters()
            
        if Field_Results == True:
            self.include_field_results()
            
        self.rep.save('my_rep.docx')
        
        print('Document Saved')
        
        
    def include_title(self):
        #Title
        self.rep.add_heading('Design Document of {}'.format(self.myID.model_parameters.name), 0)
        
        #Here to add Name, Date
        
        #Page Break
        self.rep.add_page_break()
        
    def include_contents(self):
        #Table of Contents
        self.rep.add_heading('Contents',1)
        
        #Page Break
        self.rep.add_page_break()
        
    def include_introduction(self):
        #Background
        self.rep.add_heading('Introduction')
        
        self.rep.add_paragraph('This report is a procedurally generated document for an undulator design. '
                          'This undulator is a {} device, with a period length of '
                          '{} mm, designed for a minimum magnetic gap of {} mm.\n'
                          'The key parameters for this device are listed in Table {}, and '
                          'the complete device description saved in XXX is shown in Appendix {}.'
                          ''.format(self.myID.model_parameters.type,self.myID.model_parameters.periodlength, 
                                    self.myID.model_parameters.minimumgap, self.tabnum, self.appnum))
        
        table1 = self.rep.add_table(rows = 1, cols = 2)
        table1.style = 'LightShading-Accent1'
        
        table1.cell(0,0).text = 'Parameter'
        table1.cell(0,1).text = 'Value'
        
        for param in my_ID.model_parameters.__dict__:
            new_row = table1.add_row().cells
            new_row[0].text = param
            new_row[1].text =  str(self.myID.model_parameters.__getattribute__(param))
            
        
        self.tabnum +=1
        self.appnum +=1
        
    def include_key_parameters(self):
        pass
    
    def include_field_results(self):
        #Background
        self.rep.add_heading('Magnetic Field Summary')
        
        if len(my_solution.scan_parameters.shiftrange) > 1:
            
            self.rep.add_paragraph('This device is a polarising device, and '
                                   'the peak magnetic field achievable is {:.3f} T vertical field (horizontal polarisation),'
                                   ' and {:.3f} T horizontal field (vertical polarisation). '
                                   'The corresponding K values can be seen in Table {}. '
                                   'The variation of this peak field with gap and shift is shown below in Fig. {} and Fig. {}. '
                                   'Plots of the field at minimum gap in horizontal and vertical polarisation modes are shown in Fig. {} and Fig. {}.'
                                   ''.format(max(abs(my_solution.results['Bmax'][0,0,:,0])), max(abs(my_solution.results['Bmax'][0,0,:,2])),
                                             self.tabnum, self.fignum, self.fignum+1, self.fignum +2, self.fignum + 3))
            
            #Table of Peak Field and K Values
            
            #plot of Bx and By peak with Gap
            self.mysolution.plot_Bpeak_vs_Gap('fig{}.png'.format(self.fignum))
            self.mysolution.plot_Bpeak_vs_Phase('fig{}.png'.format(self.fignum+1))
            
            sleep(14)
            
            self.rep.add_picture('fig{}.png'.format(self.fignum))
            self.rep.add_picture('fig{}.png'.format(self.fignum+1))
            
            #Plot of HH minimum Gap
            
            #Plot of VV minimum Gap
            
        
    def demo_function(self):
                
        
        rep.add_paragraph('A representative {} period model is shown in Figure {}.'.format(myID.model_parameters.periods, fignum))
        fignum+=1
        
        #Figure of device
        
        ###templates...
        #Page Break
        rep.add_page_break()
         
        # Add a heading of level 0 (Also called Title)
        rep.add_heading('Document Title, or Heading Level 0', 0)
          
        # Add a heading of level 1
        rep.add_heading('Heading level 1', 1)
          
        # Add a heading of level 2
        rep.add_heading('Heading level 2', 2)
          
        # Add a heading of level 3
        rep.add_heading('Heading level 3', 3)
          
        # Add a heading of level 4
        rep.add_heading('Heading level 4', 4)
          
        # Add a heading of level 5
        rep.add_heading('Heading level 5', 5)
          
        # Add a heading of level 6
        rep.add_heading('Heading level 6', 6)
          
        # Add a heading of level 7
        rep.add_heading('Heading level 7', 7)
          
        # Add a heading of level 8
        rep.add_heading('Heading level 8', 8)
          
        # Add a heading of level 9
        rep.add_heading('Heading level 9', 9)
          
        #
        
        para1 = rep.add_paragraph('This is my first sentence\n \mu in the report')
        
        
        rep.save('my_rep.docx')
        
        print('The Document has been saved')

    
    
def vtkeg(myID):
    #setting up argument parser. Need help, report type, input parameters, output file
    colors = vtk.vtkNamedColors()
    # Set the background color.
    bkg = map(lambda x: x / 255.0, [26, 51, 102, 255])
    colors.SetColor("BkgColor", *bkg)

    # This creates a polygonal cylinder model with eight circumferential
    # facets.
    cylinder = vtk.vtkCylinderSource()
    cylinder.SetResolution(8)

    # The mapper is responsible for pushing the geometry into the graphics
    # library. It may also do color mapping, if scalars or other
    # attributes are defined.
    cylinderMapper = vtk.vtkPolyDataMapper()
    cylinderMapper.SetInputConnection(cylinder.GetOutputPort())

    # The actor is a grouping mechanism: besides the geometry (mapper), it
    # also has a property, transformation matrix, and/or texture map.
    # Here we set its color and rotate it -22.5 degrees.
    cylinderActor = vtk.vtkActor()
    cylinderActor.SetMapper(cylinderMapper)
    cylinderActor.GetProperty().SetColor(colors.GetColor3d("Tomato"))
    cylinderActor.RotateX(30.0)
    cylinderActor.RotateY(-45.0)

    # Create the graphics structure. The renderer renders into the render
    # window. The render window interactor captures mouse events and will
    # perform appropriate camera or actor manipulation depending on the
    # nature of the events.
    ren = vtk.vtkRenderer()
    renWin = vtk.vtkRenderWindow()
    renWin.AddRenderer(ren)
    iren = vtk.vtkRenderWindowInteractor()
    iren.SetRenderWindow(renWin)

    # Add the actors to the renderer, set the background and size
    ren.AddActor(cylinderActor)
    ren.SetBackground(colors.GetColor3d("BkgColor"))
    renWin.SetSize(300, 300)
    renWin.SetWindowName('CylinderExample')

    # This allows the interactor to initalize itself. It has to be
    # called before an event loop.
    iren.Initialize()

    # We'll zoom in a little by accessing the camera and invoking a "Zoom"
    # method on it.
    ren.ResetCamera()
    ren.GetActiveCamera().Zoom(1.5)
    renWin.Render()

    # Start the event loop.
    iren.Start()
        
if __name__ == '__main__':
       
    parser = argparse.ArgumentParser()
    
    parser.add_argument("-i","--input", help = "name of input file")
    
    args = parser.parse_args()
    
    fname = args.input
    
    test_hyper_params = parameters.model_parameters(type  = "Compensated_APPLE", #Or "Plain_APPLE"
                                                    periods = 7, # Number of Periods of the APPLE Undulator
                                                    periodlength = 24, # The period length of the undulator
                                                    magnets_per_period = 4, # This number is almost exclusively 4 in undulator Halbach arrays. But it doesn't *have* to be.
                                                    
                                                    #####  APPLE Undulator Parameters  #####
                                                    rowtorowgap= 0.5, # for APPLE devices the distance between functional rows on the same jaw
                                                    end_separation = 2.5, #separation of end magnet in usual APPLE end constellation
                                                    
                                                    #####  Compensated APPLE Undulator Parameters  #####
                                                    compappleseparation = 15.0, # The gap between functional magnets and compenation magnets
                                                    
                                                    #####  Magnet Shape  #####
                                                    
                                                    nominal_fmagnet_dimensions = [20.0,0.0,20.0], # The nominal maximal magnet dimension for the functional magnets [mm]
                                                    apple_clampcut = 5.0, # The size of the square removed for clamping an APPLE magnet [mm]
                                                    apple_clampcut_non_symmetric = [5.0, 0.0, 3.0],
                                                    magnet_chamfer = [5.0,0.0,5.0], # Dimensions of chamfer for a rectangular magnet (to make it octagonal) [mm]
                                                    
                                                
                                                    #####  Compensation Magnets #####
                                                    
                                                    nominal_cmagnet_dimensions = [15.0,0.0,30.0], # dimensions of the compensation magnets [mm]
                                                    comp_magnet_chamfer = [5.0,0.0,5.0],
                                                    nominal_hcmagnet_dimensions = [15.0,0.0,30.0], # dimensions of the compensation magnets [mm]
                                                    hcomp_magnet_chamfer = [5.0,0.0,5.0],
                                                    nominal_vcmagnet_dimensions = [15.0,0.0,30.0], # dimensions of the compensation magnets [mm]
                                                    vcomp_magnet_chamfer = [5.0,0.0,5.0],
                                                    
                                                    #####  Magnet Material #####
                                                    
                                                    ksi = [.019, .06], # Permeability - anisotropic
                                                    #M = 1.21*1.344, # Block Remanence [T] Default Cryogenic Grade
                                                    M = 1.344,
                                                    Mova =  0.0, # Off Vertical Angle of Vertical type magnet blocks [degrees]
                                                    )
                                                    
    
    test_hyper_params.name = "Test_ID_Jul24"
    
    #test_hyper_params.load(fname)
    
    my_ID = id.compensatedAPPLEv2(test_hyper_params)
    
     ### Developing Model Solution ### Range of gap. rowshift and shiftmode ###
    gaprange = np.array([6,8])
    shiftrange = np.arange(-12,9.5, 12)
    shiftmoderange = ['linear','circular']
    
    #scan_parameters = parameters.scan_parameters(periodlength = test_hyper_params.periodlength, gaprange = gaprange, shiftrange = shiftrange, shiftmoderange = shiftmoderange)
    scan_parameters = parameters.scan_parameters(periodlength = test_hyper_params.periodlength, gaprange = gaprange, shiftrange = shiftrange, shiftmoderange = shiftmoderange)
    
    my_solution = af.Solution(test_hyper_params, scan_parameters,property = ['B'])
    
    my_solution.solve(property = ['B'])
    
    #vtkeg(my_ID)
    
    my_rep = Sandbox(my_ID,my_solution)
    
    my_rep.publish_report(fname = '19mm_APPLE2p5', all = True)
    
    print(1)